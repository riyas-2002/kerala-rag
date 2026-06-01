"""
Kerala RAG — Document Processing Service
Handles extraction, cleaning, chunking, deduplication, and metadata tagging
for PDF, DOCX, TXT, HTML, and Markdown files.
"""
import os
import re
import hashlib
import json
import unicodedata
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, asdict

import pdfplumber
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import markdown
import html2text
from loguru import logger

from app.core.config import settings


CATEGORY_MAP = {
    "acts_rules": "Acts & Rules",
    "licenses": "Licenses & Permits",
    "sop_guidelines": "SOPs & Guidelines",
    "forms": "Forms & Applications",
    "faqs": "FAQs",
    "central_laws": "Central Laws",
    "business_maps": "Business Type Mappings",
    "metadata": "Metadata",
}


@dataclass
class DocumentChunk:
    chunk_id: str
    text: str
    source_file: str
    source_path: str
    category: str
    page_number: Optional[int]
    chunk_index: int
    total_chunks: int
    doc_title: str
    metadata: Dict
    content_hash: str

    def to_dict(self) -> Dict:
        return asdict(self)


class DocumentProcessor:
    """
    Processes documents from the kerala_rag folder structure into
    clean, deduped, chunked text with metadata.
    """

    def __init__(self):
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
        self._seen_hashes: set = set()

    # ------------------------------------------------------------------ #
    #  Public API                                                          #
    # ------------------------------------------------------------------ #

    def process_all_documents(self) -> List[DocumentChunk]:
        """Scan the entire kerala_rag folder and process every document."""
        all_chunks: List[DocumentChunk] = []
        self._seen_hashes.clear()

        base = Path(settings.documents_path)
        supported = {".pdf", ".docx", ".txt", ".html", ".htm", ".md"}

        for category_dir in sorted(base.iterdir()):
            if not category_dir.is_dir():
                continue
            category = category_dir.name
            if category == "metadata":
                continue

            for doc_path in sorted(category_dir.rglob("*")):
                if doc_path.suffix.lower() not in supported:
                    continue
                try:
                    chunks = self.process_single_document(doc_path, category)
                    all_chunks.extend(chunks)
                    logger.info(
                        f"Processed {doc_path.name}: {len(chunks)} chunks"
                    )
                except Exception as e:
                    logger.error(f"Failed to process {doc_path}: {e}")

        logger.info(f"Total chunks produced: {len(all_chunks)}")
        return all_chunks

    def process_single_document(
        self, path: Path, category: Optional[str] = None
    ) -> List[DocumentChunk]:
        """Extract text from one file, chunk it, and return DocumentChunks."""
        if category is None:
            # Infer from parent folder name
            parent = path.parent.name
            category = parent if parent in CATEGORY_MAP else "general"

        ext = path.suffix.lower()
        pages: List[Tuple[str, Optional[int]]] = []  # (text, page_num)

        if ext == ".pdf":
            pages = self._extract_pdf(path)
        elif ext == ".docx":
            pages = self._extract_docx(path)
        elif ext in (".html", ".htm"):
            pages = self._extract_html(path)
        elif ext == ".md":
            pages = self._extract_markdown(path)
        else:  # .txt
            pages = self._extract_txt(path)

        # Load optional sidecar metadata
        meta = self._load_sidecar_metadata(path)

        # Combine all pages into one text stream but track page origins
        all_chunks: List[DocumentChunk] = []
        doc_title = meta.get("title", path.stem.replace("_", " ").title())

        # Chunk per-page to preserve page numbers
        chunk_idx = 0
        raw_chunks_by_page: List[Tuple[List[str], Optional[int]]] = []

        for page_text, page_num in pages:
            text_chunks = self._chunk_text(page_text)
            raw_chunks_by_page.append((text_chunks, page_num))

        # Flatten and count total
        flat: List[Tuple[str, Optional[int]]] = []
        for text_chunks, page_num in raw_chunks_by_page:
            for tc in text_chunks:
                flat.append((tc, page_num))

        total = len(flat)
        for raw_text, page_num in flat:
            clean = self._clean_text(raw_text)
            if len(clean) < 60:
                continue
            content_hash = hashlib.md5(clean.encode()).hexdigest()
            if content_hash in self._seen_hashes:
                continue
            self._seen_hashes.add(content_hash)

            chunk_id = f"{path.stem}_{chunk_idx}_{content_hash[:8]}"
            dc = DocumentChunk(
                chunk_id=chunk_id,
                text=clean,
                source_file=path.name,
                source_path=str(path.relative_to(settings.documents_path)),
                category=CATEGORY_MAP.get(category, category),
                page_number=page_num,
                chunk_index=chunk_idx,
                total_chunks=total,
                doc_title=doc_title,
                metadata=meta,
                content_hash=content_hash,
            )
            all_chunks.append(dc)
            chunk_idx += 1

        return all_chunks

    # ------------------------------------------------------------------ #
    #  Text Extractors                                                     #
    # ------------------------------------------------------------------ #

    def _extract_pdf(self, path: Path) -> List[Tuple[str, int]]:
        pages = []
        try:
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text() or ""
                    if len(text.strip()) < 30:
                        # Attempt OCR fallback via pytesseract
                        try:
                            import pytesseract
                            from PIL import Image
                            img = page.to_image(resolution=200).original
                            text = pytesseract.image_to_string(img, lang="eng")
                        except Exception as ocr_err:
                            logger.warning(f"OCR failed page {i} in {path.name}: {ocr_err}")
                    pages.append((text, i))
        except Exception as e:
            logger.error(f"PDF extraction error {path}: {e}")
        return pages

    def _extract_docx(self, path: Path) -> List[Tuple[str, None]]:
        try:
            doc = DocxDocument(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)
            return [("\n".join(paragraphs), None)]
        except Exception as e:
            logger.error(f"DOCX extraction error {path}: {e}")
            return []

    def _extract_html(self, path: Path) -> List[Tuple[str, None]]:
        try:
            raw = path.read_bytes()
            encoding = "utf-8"
            try:
                import chardet
                detected = chardet.detect(raw)
                encoding = detected.get("encoding", "utf-8") or "utf-8"
            except Exception:
                pass
            html_content = raw.decode(encoding, errors="replace")
            soup = BeautifulSoup(html_content, "lxml")
            # Remove nav, header, footer, script, style
            for tag in soup(["nav", "header", "footer", "script", "style", "aside"]):
                tag.decompose()
            h = html2text.HTML2Text()
            h.ignore_links = False
            h.ignore_images = True
            h.body_width = 0
            text = h.handle(str(soup))
            return [(text, None)]
        except Exception as e:
            logger.error(f"HTML extraction error {path}: {e}")
            return []

    def _extract_markdown(self, path: Path) -> List[Tuple[str, None]]:
        try:
            md_text = path.read_text(encoding="utf-8", errors="replace")
            # Convert to HTML then strip tags for clean text
            html_content = markdown.markdown(md_text)
            soup = BeautifulSoup(html_content, "lxml")
            text = soup.get_text(separator="\n")
            return [(text, None)]
        except Exception as e:
            logger.error(f"Markdown extraction error {path}: {e}")
            return []

    def _extract_txt(self, path: Path) -> List[Tuple[str, None]]:
        try:
            raw = path.read_bytes()
            encoding = "utf-8"
            try:
                import chardet
                detected = chardet.detect(raw)
                encoding = detected.get("encoding", "utf-8") or "utf-8"
            except Exception:
                pass
            text = raw.decode(encoding, errors="replace")
            return [(text, None)]
        except Exception as e:
            logger.error(f"TXT extraction error {path}: {e}")
            return []

    # ------------------------------------------------------------------ #
    #  Chunking                                                            #
    # ------------------------------------------------------------------ #

    def _chunk_text(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks of ~chunk_size tokens.
        Uses sentence-aware splitting to preserve semantic coherence.
        """
        # Rough token estimate: 1 token ≈ 4 characters
        char_size = self.chunk_size * 4
        char_overlap = self.chunk_overlap * 4

        # Split on paragraph boundaries first
        paragraphs = re.split(r"\n{2,}", text)
        chunks = []
        current = ""

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            if len(current) + len(para) <= char_size:
                current += ("\n\n" + para) if current else para
            else:
                if current:
                    chunks.append(current)
                    # Overlap: carry last part of previous chunk
                    overlap_text = current[-char_overlap:] if len(current) > char_overlap else current
                    current = overlap_text + "\n\n" + para
                else:
                    # Single paragraph larger than chunk size — split by sentence
                    sentences = re.split(r"(?<=[.!?])\s+", para)
                    sub = ""
                    for sent in sentences:
                        if len(sub) + len(sent) <= char_size:
                            sub += (" " + sent) if sub else sent
                        else:
                            if sub:
                                chunks.append(sub)
                                overlap_text = sub[-char_overlap:] if len(sub) > char_overlap else sub
                                sub = overlap_text + " " + sent
                            else:
                                chunks.append(sent)
                    if sub:
                        current = sub

        if current:
            chunks.append(current)

        return chunks

    # ------------------------------------------------------------------ #
    #  Cleaning                                                            #
    # ------------------------------------------------------------------ #

    def _clean_text(self, text: str) -> str:
        """Remove junk, normalize whitespace, strip headers/footers."""
        # Normalize unicode
        text = unicodedata.normalize("NFKC", text)

        # Remove page numbers (standalone digits on a line)
        text = re.sub(r"^\s*\d+\s*$", "", text, flags=re.MULTILINE)

        # Remove common header/footer patterns
        junk_patterns = [
            r"Page\s+\d+\s+of\s+\d+",
            r"www\.[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}",  # URLs often in footers
            r"©.*?\d{4}",
            r"All Rights Reserved",
            r"Confidential",
            r"DRAFT",
            r"_{5,}",   # underscores used as dividers
            r"-{5,}",   # dashes used as dividers
        ]
        for pattern in junk_patterns:
            text = re.sub(pattern, "", text, flags=re.IGNORECASE)

        # Collapse multiple blank lines
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Collapse multiple spaces
        text = re.sub(r" {2,}", " ", text)

        # Strip leading/trailing whitespace
        text = text.strip()

        return text

    # ------------------------------------------------------------------ #
    #  Metadata                                                            #
    # ------------------------------------------------------------------ #

    def _load_sidecar_metadata(self, doc_path: Path) -> Dict:
        """Load optional .json sidecar file next to the document."""
        sidecar = doc_path.with_suffix(".json")
        if sidecar.exists():
            try:
                return json.loads(sidecar.read_text(encoding="utf-8"))
            except Exception:
                pass
        # Also check metadata/ subfolder at same category level
        meta_dir = doc_path.parent.parent / "metadata"
        meta_file = meta_dir / (doc_path.stem + ".json")
        if meta_file.exists():
            try:
                return json.loads(meta_file.read_text(encoding="utf-8"))
            except Exception:
                pass
        return {}
